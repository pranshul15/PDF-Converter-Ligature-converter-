import os
import re
import json
import shutil
import sys
import threading
import time
from datetime import datetime
from docx import Document
import ollama

PROGRESS_FILE = "job_progress.json"
console_lock = threading.Lock()

class JobManager:
    def __init__(self, file_path, model_name="qwen3.5:2b", max_words_per_chunk=150):
        self.file_path = file_path
        self.model_name = model_name
        self.max_words_per_chunk = max_words_per_chunk
        
        # States: "IDLE", "RUNNING", "PAUSED", "STOPPED"
        self.state = "IDLE"
        self.worker_thread = None
        self.current_paragraph_tracking = 0
        self.total_paragraphs_tracking = 0

    def log(self, message):
        """
        Thread-safe console logging. Clears the current input line, 
        prints the log message, and restores the prompt cleanly.
        """
        with console_lock:
            # \r moves cursor to start, \033[K clears the line line from cursor to end
            sys.stdout.write(f"\r\033[K{message}\n")
            sys.stdout.write("> ")
            sys.stdout.flush()

    def load_progress(self):
        if os.path.exists(PROGRESS_FILE):
            try:
                with open(PROGRESS_FILE, 'r') as f:
                    return json.load(f)
            except Exception:
                return None
        return None

    def save_progress(self, current_index, total_paragraphs, backup_path):
        progress_data = {
            "file_path": self.file_path,
            "current_index": current_index,
            "total_paragraphs": total_paragraphs,
            "backup_path": backup_path,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        with open(PROGRESS_FILE, 'w') as f:
            json.dump(progress_data, f, indent=4)

    def clear_progress(self):
        if os.path.exists(PROGRESS_FILE):
            os.remove(PROGRESS_FILE)
            self.log("[Job Manager] Progress cache cleared.")

    def start(self):
        if self.state == "RUNNING":
            self.log("Job is already running!")
            return

        self.state = "RUNNING"
        self.log(f"[Job Manager] Job state updated to: {self.state}")
        self.worker_thread = threading.Thread(target=self._process_worker, daemon=True)
        self.worker_thread.start()

    def pause(self):
        if self.state == "RUNNING":
            self.state = "PAUSED"
            self.log("[Job Manager] Pause requested! Finishing current paragraph before pausing...")
        else:
            self.log("Job is not running. Cannot pause.")

    def stop(self):
        if self.state in ["RUNNING", "PAUSED"]:
            self.state = "STOPPED"
            self.log("[Job Manager] Stop requested! Halting operations and resetting progress...")
        else:
            self.log("No active job to stop.")

    def get_status(self):
        """Returns a snapshot of where the job currently stands."""
        if self.state == "RUNNING":
            return f"Status: RUNNING (Paragraph {self.current_paragraph_tracking}/{self.total_paragraphs_tracking})"
        return f"Status: {self.state}"

    def _process_worker(self):
        if not os.path.exists(self.file_path):
            self.log(f"Error: The file '{self.file_path}' does not exist.")
            self.state = "IDLE"
            return

        progress = self.load_progress()
        start_index = 1
        backup_path = None

        if progress and progress["file_path"] == self.file_path:
            start_index = progress["current_index"]
            backup_path = progress["backup_path"]
            self.log(f"[Job Manager] Resuming saved job from paragraph {start_index}...")
        else:
            backup_path = self._create_backup()
            if not backup_path:
                self.log("Failed to create backup. Aborting job.")
                self.state = "IDLE"
                return

        doc = Document(self.file_path)
        self.total_paragraphs_tracking = len(doc.paragraphs)
        changes_made = 0

        for index in range(start_index, self.total_paragraphs_tracking + 1):
            self.current_paragraph_tracking = index

            # Evaluation of State Interruptions
            if self.state == "PAUSED":
                self.log(f"[Job Manager] Job successfully PAUSED at paragraph {index}.")
                self.save_progress(index, self.total_paragraphs_tracking, backup_path)
                return 

            if self.state == "STOPPED":
                self.log("[Job Manager] Job TERMINATED completely.")
                self.clear_progress()
                return 

            paragraph = doc.paragraphs[index - 1]
            original_text = paragraph.text

            if original_text.strip():
                self.log(f"Processing paragraph {index}/{self.total_paragraphs_tracking}...")
                chunks = self._split_large_paragraph(original_text)
                corrected_chunks = []
                
                for chunk in chunks:
                    # Double check state inside chunking loops for swift execution kills
                    if self.state == "STOPPED":
                        return
                    corrected_chunks.append(self._split_combined_words(chunk))
                
                new_text = " ".join(corrected_chunks)
                
                if new_text != original_text:
                    self._merge_spaces_into_runs(paragraph, new_text)
                    changes_made += 1
                    doc.save(self.file_path)

        self.log(f"[Job Manager] Processing finished completely! {changes_made} paragraphs updated.")
        self.clear_progress()
        self.state = "IDLE"

    # --- Underlying Processing Engines ---

    def _create_backup(self):
        dir_name, file_name = os.path.split(self.file_path)
        base_name, ext = os.path.splitext(file_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(dir_name, f"{base_name}_backup_{timestamp}{ext}")
        self.log(f"[Backup] Creating baseline backup copy at: {backup_path}")
        shutil.copy2(self.file_path, backup_path)
        return backup_path

    def _split_large_paragraph(self, text):
        words = text.split()
        if len(words) <= self.max_words_per_chunk:
            return [text]
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks, current_chunk, count = [], [], 0
        for s in sentences:
            s_words = len(s.split())
            if count + s_words > self.max_words_per_chunk and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk, count = [s], s_words
            else:
                current_chunk.append(s)
                count += s_words
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        return chunks

    def _split_combined_words(self, text):
        prompt = (
            "You are a text repair assistant. Your ONLY job is to identify words that have "
            "been accidentally squashed together without a space and insert the correct space.\n"
            "Example: 'the work exceptionalteam' -> 'the work exceptional team'.\n\n"
            "CRITICAL RULES:\n"
            "1. Do NOT change style, wording, grammar, or punctuation.\n"
            "2. Do NOT fix typos unless it is a missing space between two valid words.\n"
            "3. If there are no merged words, return the original text exactly as it is.\n"
            "4. Output ONLY the repaired text. No explanations, no introduction.\n\n"
            f"Text to process:\n{text}"
        )
        try:
            response = ollama.generate(model=self.model_name, prompt=prompt, options={"temperature": 0.0})
            return response['response'].strip()
        except Exception as e:
            self.log(f"Ollama Connection Error: {e}")
            return text

    def _merge_spaces_into_runs(self, paragraph, corrected_text):
        runs = paragraph.runs
        if not runs: return
        orig_chars = [{'char': c, 'run_idx': i} for i, r in enumerate(runs) for c in r.text]
        new_run_texts = ["" for _ in runs]
        orig_ptr = 0

        for c in corrected_text:
            if orig_ptr < len(orig_chars) and c == orig_chars[orig_ptr]['char']:
                new_run_texts[orig_chars[orig_ptr]['run_idx']] += c
                orig_ptr += 1
            elif c == ' ':
                r_idx = orig_chars[orig_ptr]['run_idx'] if orig_ptr < len(orig_chars) else len(runs) - 1
                new_run_texts[r_idx] += c
            else:
                if orig_ptr < len(orig_chars):
                    new_run_texts[orig_chars[orig_ptr]['run_idx']] += c
                    orig_ptr += 1

        for run_idx, text in enumerate(new_run_texts):
            runs[run_idx].text = text


# --- Interactive CLI Loop ---
if __name__ == "__main__":
    TARGET_FILE = "The DaVinci Code.docx"
    MODEL = "gemma4:e4b"
    
    manager = JobManager(file_path=TARGET_FILE, model_name=MODEL)
    
    print("=======================================================")
    print(" Document Processor System Online                      ")
    print(" Commands: start | pause | stop | status | exit        ")
    print("=======================================================")
    
    # Initialize the first input layout line
    sys.stdout.write("> ")
    sys.stdout.flush()

    while True:
        try:
            command = input().strip().lower()
        except (KeyboardInterrupt, EOFError):
            manager.stop()
            break
        
        if command == "start":
            manager.start()
        elif command == "pause":
            manager.pause()
        elif command == "stop":
            manager.stop()
        elif command == "status":
            manager.log(manager.get_status())
        elif command == "exit":
            manager.stop()
            print("Exiting application safely.")
            break
        elif command == "":
            # If the user just presses Enter, rewrite the clean arrow prompt
            sys.stdout.write("> ")
            sys.stdout.flush()
        else:
            manager.log(f"Unknown command: '{command}'. Available: start, pause, stop, status, exit")