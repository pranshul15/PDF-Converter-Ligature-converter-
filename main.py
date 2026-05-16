import subprocess
import os
import argparse
from pathlib import Path
from pdf2docx import Converter
from docx import Document

def rebuild_pdf(input_file, output_file):
    print(f"--- Step 1: Rebuilding {input_file} with Ghostscript ---")
    
    # Common Ghostscript executable names
    gs_cmds = ["gs", "gswin64c", "gswin32c"]
    
    for cmd in gs_cmds:
        args = [
            cmd, 
            "-dNOPAUSE", "-dBATCH", "-dSAFER",
            "-sDEVICE=pdfwrite",
            f"-sOutputFile={output_file}",
            input_file
        ]
        try:
            # We use subprocess.run and catch FileNotFoundError if the command isn't found
            subprocess.run(args, check=True, capture_output=True, text=True)
            print("PDF rebuilt successfully.")
            return True
        except FileNotFoundError:
            # This command wasn't found, try the next one
            continue
        except subprocess.CalledProcessError as e:
            print(f"Error during Ghostscript execution ({cmd}):")
            print(e.stderr)
            return False

    # If we loop through all and none work
    print("\n[!] Error: Ghostscript not found.")
    print("Ghostscript is required to 'rebuild' the PDF and fix ligature issues.")
    print("1. Install Ghostscript: https://ghostscript.com/releases/gsdnld.html")
    print("2. Ensure 'gs' or 'gswin64c' is in your system PATH.")
    print("3. Alternatively, run this tool via Docker (recommended).")
    return False

def convert_to_docx(pdf_file, docx_file):
    print(f"--- Step 2: Converting {pdf_file} to Word ---")
    try:
        cv = Converter(pdf_file)
        cv.convert(docx_file, multi_processing=True)
        cv.close()
        print(f"Successfully created {docx_file}")
        return True
    except Exception as e:
        print(f"Error during conversion: {e}")
        return False

def remove_section_breaks_with_space(docx_file):
    print(f"--- Step 3: Removing section breaks from {docx_file} ---")
    try:
        doc = Document(docx_file)
        for paragraph in doc.paragraphs:
            p_element = paragraph._p
            p_pr = p_element.pPr
            
            if p_pr is not None:
                sect_prs = p_pr.xpath('./w:sectPr')
                if sect_prs:
                    for sect_pr in sect_prs:
                        p_pr.remove(sect_pr)
                    paragraph.add_run(' ')
        
        doc.save(docx_file)
        print(f"Done! Section breaks replaced with spaces in {docx_file}")
    except Exception as e:
        print(f"Error removing section breaks: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert PDF to DOCX with ligature and section break fixes.")
    parser.add_argument("-i", "--input", default="book.pdf", help="Input PDF filename (default: book.pdf)")
    parser.add_argument("-o", "--output", help="Output DOCX filename (default: [input]_final.docx)")
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found in the directory.")
        exit(1)

    # Derive filenames
    input_stem = input_path.stem
    rebuilt_pdf = f"{input_stem}_rebuilt.pdf"
    output_docx = args.output if args.output else f"{input_stem}_final.docx"

    # Run Sequence
    if rebuild_pdf(str(input_path), rebuilt_pdf):
        if convert_to_docx(rebuilt_pdf, output_docx):
            remove_section_breaks_with_space(output_docx)
            # Optional: Clean up intermediate rebuilt PDF
            if os.path.exists(rebuilt_pdf):
                os.remove(rebuilt_pdf)